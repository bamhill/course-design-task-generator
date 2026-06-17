# -*- coding: utf-8 -*-
"""生成《计算机网络课程设计》课程设计任务书（初版）.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx.oxml import OxmlElement
def set_cell(cell, text, bold=False, size=Pt(10.5)):
    cell.text = ''
    # 垂直居中
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), 'center')
    tcPr.append(tcVAlign)
    # 段落居中
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = size; run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_p(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = Pt(11); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_sipin(text):
    p = doc.add_paragraph()
    r1 = p.add_run('💡 思政融入提示：')
    r1.font.name = '宋体'; r1.font.size = Pt(10); r1.bold = True
    r1.font.color.rgb = RGBColor(43, 95, 138)
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r2 = p.add_run(text)
    r2.font.name = '宋体'; r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(43, 95, 138)
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ═══ 标题 ═══
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('《计算机网络课程设计》\n课程设计任务书')
r.font.name = '黑体'; r.font.size = Pt(18); r.bold = True
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('课程名称', '计算机网络课程设计'),
    ('课程编号', '04040358b'),
    ('学分学时', '1 学分，1 周'),
    ('适用专业', '信息管理与信息系统 · 2024 级'),
    ('学生人数', '70 人'),
]
for i, (k, v) in enumerate(info_data):
    set_cell(info_table.cell(i, 0), k, bold=True)
    set_cell(info_table.cell(i, 1), v)
doc.add_paragraph()

# ═══ 一、课程设计目标 ═══
add_h('一、课程设计目标')
add_p('本课程设计旨在使学生修完《计算机网络》理论课后，开展的实践性环节。通过计算机网络设计、配置和调试等综合实践，使学生对计算机网络的核心知识——网络方案设计、VLAN 划分、路由配置、应用服务配置与安全配置——得到全面巩固和提高。', indent=True)
add_p('通过本课程设计，学生应达到以下目标：', indent=True)
add_p('目标1（设计能力）：能分析、设计小型计算机网络，制定合适的网络配置方案。', indent=True)
add_p('目标2（配置能力）：能根据方案在模拟器中完成网络的基本连通调试，包括交换机VLAN划分、三层交换机配置和路由配置；采用DHCP进行IP自动配置；通过VPN技术实现远程接入内部网。', indent=True)
add_p('目标3（应用与安全）：能根据方案实现应用层网络服务配置（Web、DNS、电子邮件、FTP等）以及相关安全配置（防DHCP欺骗、防火墙等）。', indent=True)

# ═══ 二、设计任务 ═══
add_h('二、设计任务')

add_h('任务一：小型网络解决方案设计', level=2)
add_p('支撑目标：课程目标1', bold=True)

add_h('阶段1：需求分析与总体方案设计', level=3)
add_p('任务内容：', bold=True)
add_p('1. 根据学生专业知识背景选择一个实验题目，对相应的小型网络进行需求分析。', indent=True)
add_p('2. 完成该小型网络的总体方案设计，确定网络逻辑拓扑结构和所采用的网络技术。', indent=True)
add_p('3. 确定主要设备的性能指标，完成设备选型和经费预算（选做）。', indent=True)
add_p('交付物：网络设计方案书（含需求分析、逻辑拓扑图、设备选型表）', bold=True)
add_p('学习要求：', bold=True)
add_p('1. 理解网络方案设计的一般流程，掌握网络方案设计的基本方法。', indent=True)
add_p('2. 能够结合实际需求进行方案设计，了解进行网络预算的基本方法。', indent=True)
add_p('3. 完成网络设计方案的编写，掌握网络逻辑拓扑结构的绘制方法。', indent=True)
add_sipin('在需求分析阶段强调网络布线、设备选型要符合国家标准，树立"一切从工程实际出发"的设计理念。引导学生查阅相关国家标准、配置资料——网络方案不是写论文，而是要从实际出发。（融入时机：讲解）')

add_h('任务二：局域网配置与实现', level=2)
add_p('支撑目标：课程目标2、3', bold=True)

add_h('阶段2：网络基础搭建与配置', level=3)
add_p('任务内容：', bold=True)
add_p('1. 在模拟器中完成网络的基本搭建以及相关配置。', indent=True)
add_p('2. 交换机VLAN划分、三层交换机配置以及路由配置。', indent=True)
add_p('3. 采用DHCP进行IP自动配置并进行调试。', indent=True)
add_p('4. 通过VPN技术实现远程接入内部网。', indent=True)
add_p('交付物：网络配置文件（交换机/路由器配置命令 + 网络拓扑截图）', bold=True)
add_p('学习要求：', bold=True)
add_p('1. 理解网络通信配置的基本原理和配置方法。', indent=True)
add_p('2. 掌握不同类型VPN以及相应的配置方法。', indent=True)
add_sipin('配置交换机、路由器时强调每一行命令都有对应的工程含义，不能"照着教程敲"。用"配置错了全网瘫痪"的极端案例提醒学生网络工程师的社会责任——你配的网络可能是一家公司、一家医院的通信命脉。（融入时机：实操/检查）')

add_h('阶段3：应用服务与安全配置', level=3)
add_p('任务内容：', bold=True)
add_p('1. 根据方案实现应用层网络服务配置（Web服务、DNS服务、电子邮件服务、FTP服务等）。', indent=True)
add_p('2. 根据方案进行相关安全配置，包括防DHCP欺骗、防火墙等。', indent=True)
add_p('交付物：应用服务配置文档 + 安全配置截图', bold=True)
add_p('学习要求：', bold=True)
add_p('1. 能够抓取各种不同协议类型的报文，并分析实验过程中的数据流。', indent=True)
add_p('2. 认真按时完成课程设计报告。', indent=True)
add_sipin('在讲我国没有根域名服务器时，延伸讨论网络安全自主可控的重要性——DNS是互联网的"电话簿"，没有根服务器意味着底层依赖他人。我国已建立镜像根服务器，IPv6计划中全球25台根服务器我国有4台。以此鼓励学生扎实学习，在科技创新中做出贡献。（融入时机：讲解/讨论）')

add_h('任务三：汇报答辩', level=2)
add_p('支撑目标：课程目标1-3', bold=True)

add_h('阶段4：答辩准备与汇报', level=3)
add_p('任务内容：', bold=True)
add_p('1. 网络整体配置检查。', indent=True)
add_p('2. 准备答辩PPT。', indent=True)
add_p('3. 小组答辩。', indent=True)
add_p('交付物：答辩PPT', bold=True)
add_p('学习要求：', bold=True)
add_p('1. 掌握计算机网络设计及配置的基本方法。', indent=True)
add_p('2. 熟悉计算机网络实践过程，答辩过程中能清楚流利地回答教师的问题。', indent=True)
add_sipin('答辩环节重点考查各成员负责的模块衔接情况——"你的配置如果错了，对全组有什么影响？"培养学生工程协作中的责任意识。（融入时机：答辩）')

# ═══ 三、进度安排 ═══
add_h('三、进度安排')
sch_table = doc.add_table(rows=6, cols=3)
sch_table.style = 'Table Grid'
sch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sch_headers = ['天数', '阶段', '主要任务']
sch_data = [
    ('第1天', '需求分析与方案设计', '选题、需求分析、拓扑设计、设备选型'),
    ('第2天', '网络基础搭建', '交换机VLAN划分、三层交换机配置、路由配置'),
    ('第3天', 'DHCP+VPN配置', 'IP自动配置、VPN远程接入调试'),
    ('第4天', '应用服务与安全配置', 'Web/DNS/邮件/FTP配置、防火墙、防DHCP欺骗'),
    ('第5天', '报告撰写+答辩准备', '课程设计报告、答辩PPT'),
]
for j, h in enumerate(sch_headers):
    set_cell(sch_table.cell(0, j), h, bold=True)
for i, row in enumerate(sch_data):
    for j, v in enumerate(row):
        set_cell(sch_table.cell(i+1, j), v)
doc.add_paragraph()

# ═══ 四、思政融合要求 ═══
add_h('四、思政融合要求')
add_p('思政元素不单独设章节，融入在各任务阶段的教学要求中。每个阶段1-2句，自然融入，不硬贴标签。', indent=True)
sipin_table = doc.add_table(rows=5, cols=3)
sipin_table.style = 'Table Grid'
sipin_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sipin_h = ['阶段', '思政融入点', '融入方式']
sipin_d = [
    ('需求分析与方案设计', '工程标准意识：网络方案从实际出发，符合国家标准', '讲解'),
    ('交换机/路由配置', '职业责任：配置命令有工程含义，失误可能导致全网瘫痪', '实操/检查'),
    ('DNS/应用服务配置', '技术自主：我国根域名服务器现状与科技创新', '讲解/讨论'),
    ('汇报答辩', '协作责任：模块衔接意识', '答辩'),
]
for j, h in enumerate(sipin_h):
    set_cell(sipin_table.cell(0, j), h, bold=True)
for i, row in enumerate(sipin_d):
    for j, v in enumerate(row):
        set_cell(sipin_table.cell(i+1, j), v)
doc.add_paragraph()

# ═══ 五、考核标准 ═══
add_h('五、考核标准')
exam_table = doc.add_table(rows=4, cols=3)
exam_table.style = 'Table Grid'
exam_table.alignment = WD_TABLE_ALIGNMENT.CENTER
exam_h = ['考核项目', '权重', '考查点']
exam_d = [
    ('平时考核', '40%', '网络设计方案、局域网搭建、IP配置与调试、VPN配置、应用与安全配置的完成情况'),
    ('答辩考核', '20%', '设计与配置合理性和准确性、团队协作意识、问题发现与解决能力'),
    ('课程设计报告考核', '40%', '报告内容的正确性、结构的完整性和格式的规范性'),
]
for j, h in enumerate(exam_h):
    set_cell(exam_table.cell(0, j), h, bold=True)
for i, row in enumerate(exam_d):
    for j, v in enumerate(row):
        set_cell(exam_table.cell(i+1, j), v)
doc.add_paragraph()

# ═══ 六、交付物清单 ═══
add_h('六、交付物清单')
del_table = doc.add_table(rows=7, cols=3)
del_table.style = 'Table Grid'
del_table.alignment = WD_TABLE_ALIGNMENT.CENTER
del_h = ['序号', '交付物', '提交形式']
del_d = [
    ('1', '网络设计方案书（含需求分析、逻辑拓扑图、设备选型）', 'Word / PDF'),
    ('2', '网络配置文件（交换机/路由器配置命令）', '.txt / .pkt'),
    ('3', '应用服务配置截图+安全配置截图', 'Word / PDF'),
    ('4', '网络拓扑搭建截图（含VLAN/路由/DHCP/VPN）', 'Word / PDF'),
    ('5', '课程设计报告（完整版）', 'Word / PDF'),
    ('6', '答辩PPT', 'PPT'),
]
for j, h in enumerate(del_h):
    set_cell(del_table.cell(0, j), h, bold=True)
for i, row in enumerate(del_d):
    for j, v in enumerate(row):
        set_cell(del_table.cell(i+1, j), v)
doc.add_paragraph()

# ═══ 七、参考资料 ═══
add_h('七、参考资料')
add_p('教材：', bold=True)
add_p('《计算机网络》（第8版），谢希仁主编，电子工业出版社，2021.06', indent=True)
add_p('参考教材：', bold=True)
add_p('1. 《深入浅出计算机网络（微课视频版）》，高军等主编，清华大学出版社，2022.08', indent=True)
add_p('2. 《计算机网络 - 自顶向下方法（原书第8版）》，James F. Kurose等，机械工业出版社，2022.01', indent=True)
add_p('3. 《网络技术基础与计算思维》，沈鑫钊等编著，清华大学出版社，2016.03', indent=True)
add_p('在线资源及AI工具：', bold=True)
add_p('1. 中国大学MOOC：计算机网络（哈工大） https://www.icourse163.org/course/HIT-154005', indent=True)
add_p('2. 华为ICT学院网络技术官方学习平台 https://e.huawei.com/cn/training/ict-academy', indent=True)
add_p('3. Cisco Packet Tracer AI Assistant —— 网络仿真+AI辅助排错', indent=True)
add_p('4. Wireshark AI Inspector 插件 —— 协议包智能解析', indent=True)
add_p('5. 通义灵码 —— AI编程助手，辅助Socket编程', indent=True)

# 保存
out_dir = os.path.dirname(os.path.abspath(__file__))
path = os.path.join(out_dir, '示例_计网课程设计任务书_初版.docx')
doc.save(path)
print(f'DONE: {path}')
