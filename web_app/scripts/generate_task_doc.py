# -*- coding: utf-8 -*-
"""生成《数据库原理与应用（SQL-Server）实践》课程设计任务书（初版）.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys, io

doc = Document()

# ─── 全局样式 ───
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def set_cell_font(cell, text, bold=False, size=Pt(10.5)):
    cell.text = ''
    # 垂直居中
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), 'center')
    tcPr.append(tcVAlign)
    # 水平居中
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(16)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_sipin(text):
    """思政提示段落"""
    p = doc.add_paragraph()
    run_label = p.add_run('💡 思政融入提示：')
    run_label.font.name = '宋体'
    run_label.font.size = Pt(10)
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(43, 95, 138)
    run_label.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_text = p.add_run(text)
    run_text.font.name = '宋体'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(43, 95, 138)
    run_text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5

# ════════════════════════════════════════
# 标题
# ════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《数据库原理与应用（SQL-Server）实践》\n课程设计任务书')
run.font.name = '黑体'
run.font.size = Pt(18)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 基本信息表
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ('课程名称', '数据库原理与应用（SQL-Server）实践'),
    ('课程编号', '04040313b'),
    ('学分学时', '1 学分，1 周（16 学时）'),
    ('适用专业', '信息管理与信息系统'),
    ('先修课程', '计算机程序设计语言（Python）、数据库原理与应用'),
]
for i, (k, v) in enumerate(info_data):
    set_cell_font(info_table.cell(i, 0), k, bold=True)
    set_cell_font(info_table.cell(i, 1), v)

doc.add_paragraph()

# ════════════════════════════════════════
# 一、课程设计目标
# ════════════════════════════════════════
add_heading_styled('一、课程设计目标', level=1)

add_para('本课程设计旨在使学生修完《数据库原理与应用》理论课后，开展的实践性环节。通过模拟企业物资采购和库存管理真实业务场景，结合虚拟实景实验教学系统，让学生将数据库原理、设计理论及SQL语言等理论知识应用于实际操作，通过"需求分析—概念结构设计—逻辑结构设计—数据库实现—功能调试维护"的完整实践流程，巩固理论基础，提升数据库设计、开发及问题解决的实践能力。', indent=True)

add_para('通过本课程设计，学生应达到以下目标：', indent=True)

add_para('目标1（知识应用）：理解数据库原理、数据库设计理论及关系数据库核心理论（数据依赖理论、规范化理论、模式分级理论等），掌握根据实际业务需求进行数据库系统设计的完整方法。', indent=True)

add_para('目标2（实践能力）：理解数据库系统实现的核心逻辑，掌握根据功能要求运用SQL语言实现数据库系统及解决实际业务问题的方法，提升相关实践与协作能力。', indent=True)

# ════════════════════════════════════════
# 二、设计任务
# ════════════════════════════════════════
add_heading_styled('二、设计任务', level=1)

# --- 任务一 ---
add_heading_styled('任务一：物资采购和库存管理数据库设计', level=2)
add_para('支撑目标：课程目标1', bold=True)

add_heading_styled('阶段1：需求分析', level=3)
add_para('任务内容：', bold=True)
add_para('1. 进入"企业管理虚拟实景实验教学系统"，完成物资采购计划、物资采购管理、物资入出库管理、经营合同管理等实验。', indent=True)
add_para('2. 通过实验模拟"跟班作业"的方式进行调研，获取并分析用户需求。', indent=True)
add_para('3. 在调查研究的基础上进行需求分析，绘制数据流图。', indent=True)

add_para('交付物：业务流程分析报告（含业务流程图、数据单据截图）', bold=True)
add_para('学习要求：', bold=True)
add_para('1. 了解"采购和库存管理"的完整业务流程，完成业务流程分析，在"业务流程分析"部分附上相关业务对应的数据单据截图，并整理相关数据。', indent=True)
add_para('2. 掌握需求分析的基本方法。', indent=True)
add_sipin('在调研采购流程时，提醒学生注意供应商信息、合同金额等业务数据的敏感性，引导学生讨论数据保密对企业的重要性。（融入时机：导入/讨论）')

add_heading_styled('阶段2：概念结构设计', level=3)
add_para('任务内容：', bold=True)
add_para('1. 在需求分析的基础上进行抽象建模，设计局部E-R模型。', indent=True)
add_para('2. 将局部E-R模型进行合并，解决冲突，形成总体概念E-R模型，完成概念结构设计。', indent=True)
add_para('交付物：E-R图（局部→全局）、需求分析说明书（含数据流图、数据字典）', bold=True)
add_para('学习要求：', bold=True)
add_para('1. 了解"企业管理虚拟实景实验教学系统"中"采购和库存管理"的完整业务流程。', indent=True)
add_para('2. 掌握采购管理系统数据库的概念结构、逻辑结构的设计方法。', indent=True)
add_sipin('用E-R图展示采购-库存-财务的数据流转，强调"局部服从整体"的系统设计思想——数据库设计也要站在企业全局视角看问题。（融入时机：讲解）')

add_heading_styled('阶段3：逻辑结构设计', level=3)
add_para('任务内容：', bold=True)
add_para('1. 将概念结构转换为关系数据模型。', indent=True)
add_para('2. 以规范化理论为指导，进行数据模型的优化，完成数据的关系模式设计。', indent=True)
add_para('交付物：逻辑结构设计报告（含关系模式、规范化分析）', bold=True)
add_para('学习要求：', bold=True)
add_para('1. 掌握E-R图向关系模式的转换规则。', indent=True)
add_para('2. 理解数据依赖理论，掌握规范化的方法与步骤。', indent=True)
add_sipin('规范化理论要求精确的数据依赖分析——模式分解不彻底会导致数据冗余和更新异常。类比工程领域的"偷工减料"：看似省事实则埋雷。（融入时机：讲解）')

# --- 任务二 ---
add_heading_styled('任务二：物资采购和库存管理数据库实现', level=2)
add_para('支撑目标：课程目标2', bold=True)

add_heading_styled('阶段4：数据库开发', level=3)
add_para('任务内容：', bold=True)
add_para('1. 创建数据库、创建数据表、设置表字段等操作。', indent=True)
add_para('2. 创建视图，使用SELECT语句查询视图中的数据。', indent=True)
add_para('3. 触发器创建，用于加强数据的完整性约束和业务规则。', indent=True)
add_para('交付物：数据库实现脚本（.sql文件）', bold=True)
add_para('学习要求：', bold=True)
add_para('1. 掌握功能模块设计的核心思路，掌握功能模块的设计与模拟实现方法。', indent=True)
add_para('2. 掌握数据库调试技能，加深对数据完整性的认知。', indent=True)
add_sipin('每条SQL写功能注释、规范命名。代码的"可读性就是可维护性"——这既是专业素养，也是对自己和后续维护者负责的态度。（融入时机：实操/检查）')

add_heading_styled('阶段5：功能调试与维护', level=3)
add_para('任务内容：', bold=True)
add_para('1. 根据需求设计应用功能，确定相关功能名称和模块名称。', indent=True)
add_para('2. 采用存储过程、函数、T-SQL或简单SQL语句等方式实现数据库系统功能。', indent=True)
add_para('3. 对数据库应用进行调试，对数据库中已入库数据进行转储和恢复等。', indent=True)
add_para('交付物：功能模块实现脚本（.sql文件）', bold=True)
add_para('学习要求：', bold=True)
add_para('1. 了解功能模块设计的核心思路，掌握功能模块的设计与模拟实现方法。', indent=True)
add_para('2. 掌握数据库调试技能，进一步理解数据完整性的核心内涵。', indent=True)
add_sipin('调试时有意识地去掉外键约束，让学生观察数据不一致的后果，讨论数据完整性对企业运营的意义——数据正确是企业的底线。（融入时机：检查）')

# ════════════════════════════════════════
# 三、进度安排
# ════════════════════════════════════════
add_heading_styled('三、进度安排', level=1)

schedule_table = doc.add_table(rows=7, cols=3)
schedule_table.style = 'Table Grid'
schedule_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['天数', '阶段', '主要任务']
schedule_data = [
    ('第1天', '需求分析', '实验操作、业务流程分析、数据流图'),
    ('第2天', '概念结构设计', '局部E-R→全局E-R'),
    ('第3天', '逻辑结构设计', '关系模式转换、规范化分析'),
    ('第4天', '数据库开发', '建库/建表/视图/触发器'),
    ('第5天', '功能调试+报告', '存储过程、调试、撰写设计报告'),
    ('第6天', '答辩', 'PPT准备+小组答辩'),
]

for j, h in enumerate(headers):
    set_cell_font(schedule_table.cell(0, j), h, bold=True)
for i, row in enumerate(schedule_data):
    for j, v in enumerate(row):
        set_cell_font(schedule_table.cell(i+1, j), v)

doc.add_paragraph()

# ════════════════════════════════════════
# 四、思政融合要求
# ════════════════════════════════════════
add_heading_styled('四、思政融合要求', level=1)
add_para('思政元素不单独设章节，融入在各任务阶段的教学要求中。每个阶段1-2句，自然融入，不硬贴标签。', indent=True)

sipin_table = doc.add_table(rows=7, cols=3)
sipin_table.style = 'Table Grid'
sipin_table.alignment = WD_TABLE_ALIGNMENT.CENTER

sipin_headers = ['阶段', '思政融入点', '融入方式']
sipin_data = [
    ('需求分析', '数据安全意识：业务数据是企业的核心资产', '导入/讨论'),
    ('概念结构设计', '系统思维：全局视角设计数据流转', '讲解'),
    ('逻辑结构设计', '工程严谨：规范化不是走形式', '讲解/作业'),
    ('数据库开发', '工匠精神：代码规范与注释', '实操/检查'),
    ('功能调试', '数据完整性是底线', '检查'),
    ('团队协作', '信息共享与责任意识', '答辩'),
]
for j, h in enumerate(sipin_headers):
    set_cell_font(sipin_table.cell(0, j), h, bold=True)
for i, row in enumerate(sipin_data):
    for j, v in enumerate(row):
        set_cell_font(sipin_table.cell(i+1, j), v)

doc.add_paragraph()

# ════════════════════════════════════════
# 五、考核标准
# ════════════════════════════════════════
add_heading_styled('五、考核标准', level=1)

exam_table = doc.add_table(rows=4, cols=3)
exam_table.style = 'Table Grid'
exam_table.alignment = WD_TABLE_ALIGNMENT.CENTER

exam_headers = ['考核项目', '权重', '考查点']
exam_data = [
    ('平时考核', '20%', '数据库设计和实施方案的完成情况、出勤、进度'),
    ('设计报告考核', '60%', '需求与功能匹配性、实现合理性、SQL规范性、结果正确性'),
    ('答辩考核', '20%', '回答问题正确性、团队协作意识、发现与解决问题能力'),
]
for j, h in enumerate(exam_headers):
    set_cell_font(exam_table.cell(0, j), h, bold=True)
for i, row in enumerate(exam_data):
    for j, v in enumerate(row):
        set_cell_font(exam_table.cell(i+1, j), v)

doc.add_paragraph()

# ════════════════════════════════════════
# 六、交付物清单
# ════════════════════════════════════════
add_heading_styled('六、交付物清单', level=1)

deliver_table = doc.add_table(rows=9, cols=3)
deliver_table.style = 'Table Grid'
deliver_table.alignment = WD_TABLE_ALIGNMENT.CENTER

deliver_headers = ['序号', '交付物', '提交形式']
deliver_data = [
    ('1', '业务流程分析报告（含数据单据截图）', 'Word / PDF'),
    ('2', '需求分析说明书（含数据流图、数据字典）', 'Word / PDF'),
    ('3', 'E-R图（局部→全局）', 'Visio / 手绘扫描'),
    ('4', '逻辑结构设计（关系模式、规范化分析）', 'Word / PDF'),
    ('5', '数据库实现脚本（建库/建表/视图/触发器）', '.sql文件'),
    ('6', '功能模块实现（存储过程、函数、T-SQL）', '.sql文件'),
    ('7', '课程设计报告（完整版）', 'Word / PDF'),
    ('8', '答辩PPT', 'PPT'),
]
for j, h in enumerate(deliver_headers):
    set_cell_font(deliver_table.cell(0, j), h, bold=True)
for i, row in enumerate(deliver_data):
    for j, v in enumerate(row):
        set_cell_font(deliver_table.cell(i+1, j), v)

doc.add_paragraph()

# ════════════════════════════════════════
# 七、参考资料
# ════════════════════════════════════════
add_heading_styled('七、参考资料', level=1)

add_para('教材：', bold=True)
add_para('《数据库系统概论（第6版）》，王珊、萨师煊，高等教育出版社，2024.11', indent=True)

add_para('参考教材：', bold=True)
add_para('1. 《数据库系统概念》，亚伯拉罕·西尔伯沙茨等，机械工业出版社，2021.06', indent=True)
add_para('2. 《数据库系统原理及应用教程（第5版）》，苗雪兰等，机械工业出版社，2020.4', indent=True)
add_para('3. 《数据库系统及应用（第5版）》，崔巍等，高等教育出版社，2025.8', indent=True)

add_para('在线资源及AI工具：', bold=True)
add_para('1. 中国大学MOOC：数据库及应用 https://www.icourse163.org/course/ZUEL-1206120803', indent=True)
add_para('2. 中国大学MOOC：数据库系统原理 https://www.icourse163.org/course/BNU-1002842007', indent=True)
add_para('3. AI编程工具：GitHub Copilot https://copilot.github.com/', indent=True)
add_para('4. AI编程工具：Claude Code https://claude.com/product/claude-code', indent=True)

add_para('任课教师提供的其他学习资料（在课程设计启动时补充）', indent=True)

# ════════════════════════════════════════
# 保存
# ════════════════════════════════════════
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, '示例_课程设计任务书_初版.docx')
doc.save(output_path)
print('DONE: ' + output_path)
