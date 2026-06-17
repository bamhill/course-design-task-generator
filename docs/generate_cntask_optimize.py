# -*- coding: utf-8 -*-
"""生成计网课程设计任务书 优化总结报告 + 对比表 + 定稿"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell(cell, text, bold=False, size=Pt(10.5)):
    cell.text = ''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign'); v.set(qn('w:val'), 'center'); tcPr.append(v)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0)
    run = p.add_run(text); run.font.name = '宋体'; run.font.size = size; run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_h(text, level=1, which=None):
    d = which if which else doc
    h = d.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_p(text, bold=False, indent=False, which=None):
    d = which if which else doc
    p = d.add_paragraph()
    run = p.add_run(text); run.font.name = '宋体'; run.font.size = Pt(11); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5

# ════════════════════════════════════════
# 优化总结报告
# ════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']; style.font.name = '宋体'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); style.paragraph_format.line_spacing = 1.5

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('《计算机网络课程设计》\n课程设计任务书 · 优化总结报告')
r.font.name = '黑体'; r.font.size = Pt(18); r.bold = True
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('信息管理与信息系统 · 2024 级 · 70 人')
r2.font.name = '宋体'; r2.font.size = Pt(12)
r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.add_paragraph()

# 一、优化概览
add_h('一、优化概览')
add_p('本报告记录了计网课程设计任务书初版在教师审阅后的优化过程。教师共提出 3 条优化需求，涉及交付物补充、报告格式要求和预算细化。以下逐条说明。', indent=True)

# 二、优化条目
add_h('二、优化条目详情', level=1)

# 条目1
add_h('优化条目 #1：交付物增加网络拓扑截图', level=2)
add_p('原内容：交付物清单中没有网络拓扑截图项。', indent=True)
add_p('优化后：在交付物清单中增加"网络拓扑图（含设备标识和IP地址分配）"，提交形式为 Visio / 绘图截图。', indent=True)
add_p('优化原因：网络拓扑是设计方案的核心产出之一，教师需要在交付物中看到最终的网络结构图，便于检查方案合理性和配线准确性。', indent=True)
add_p('涉及章节：§六 交付物清单', indent=True)
add_p('', indent=False)

# 条目2
add_h('优化条目 #2：要求在报告中撰写"遇到的问题及解决方案"', level=2)
add_p('原内容：课程设计报告的要求中，没有明确要求记录问题和解决过程。', indent=True)
add_p('优化后：在§三 进度安排的第5天任务中增加"撰写《课程设计报告》（含遇到的问题及解决方案）"，同时在交付物清单中注明报告必须包含此章节。', indent=True)
add_p('优化原因：教师认为记录和解决问题的过程是课设最重要的学习产出之一，不能只交一份"漂亮的"配置结果，要能看到学生的思考过程。', indent=True)
add_p('涉及章节：§三 进度安排、§六 交付物清单', indent=True)
add_p('', indent=False)

# 条目3
add_h('优化条目 #3：设计方案中明确预算要求', level=2)
add_p('原内容：设计方案中的设备选型和经费预算标注为"选做"。', indent=True)
add_p('优化后：经费预算改为必做项，要求在方案书中明确列出设备清单、型号、数量、单价和总价。', indent=True)
add_p('优化原因：教师认为预算是工程方案的重要组成部分，"选做"会导致大多数学生跳过，达不到工程实践训练的目的。', indent=True)
add_p('涉及章节：§二 阶段1 任务内容', indent=True)
add_p('', indent=False)

# 三、改动对比表
add_h('三、改动对比表', level=1)
comp_table = doc.add_table(rows=4, cols=4)
comp_table.style = 'Table Grid'; comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
comp_h = ['#', '改动位置', '原内容', '改后内容']
comp_d = [
    ('1', '§六 交付物清单', '无网络拓扑截图项', '增加"网络拓扑图（含设备标识和IP分配）"'),
    ('2', '§三+§六', '报告不要求记录问题', '增加"遇到的问题及解决方案"章节'),
    ('3', '§二 阶段1 任务内容', '经费预算标注"选做"', '经费预算改为必做，含设备单价和总价'),
]
for j, h in enumerate(comp_h):
    set_cell(comp_table.cell(0, j), h, bold=True)
for i, row in enumerate(comp_d):
    for j, v in enumerate(row):
        set_cell(comp_table.cell(i+1, j), v)
doc.add_paragraph()

# 四、优化结果
add_h('四、优化结果')
add_p('以上 3 条优化需求已全部融入任务书定稿。', indent=True)

doc.add_paragraph()
add_h('附：课程设计任务书（定稿）')
add_p('优化后的任务书定稿文件已生成，包含以上全部优化内容。', indent=True)

out_dir = os.path.dirname(os.path.abspath(__file__))
doc.save(os.path.join(out_dir, '示例_计网优化总结报告+对比表.docx'))
print('DONE1')

# ════════════════════════════════════════
# 任务书定稿
# ════════════════════════════════════════
doc2 = Document()
style2 = doc2.styles['Normal']; style2.font.name = '宋体'; style2.font.size = Pt(11)
style2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); style2.paragraph_format.line_spacing = 1.5

title2 = doc2.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title2.add_run('《计算机网络课程设计》\n课程设计任务书（定稿）')
r.font.name = '黑体'; r.font.size = Pt(18); r.bold = True
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

info = doc2.add_table(rows=5, cols=2); info.style = 'Table Grid'; info.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate([
    ('课程名称', '计算机网络课程设计'), ('课程编号', '04040358b'),
    ('学分学时', '1 学分，1 周'), ('适用专业', '信息管理与信息系统 · 2024 级 · 70 人'),
    ('优化说明', '已融入 3 条教师优化需求（v1.1）'),
]):
    set_cell(info.cell(i, 0), k, bold=True)
    set_cell(info.cell(i, 1), v)
doc2.add_paragraph()

add_h('一、课程设计目标', which=doc2)
add_p('（同初版，不变）', indent=True, which=doc2)
add_p('目标1（设计能力）：能分析、设计小型计算机网络，制定合适的网络配置方案。', indent=True, which=doc2)
add_p('目标2（配置能力）：能完成交换机VLAN划分、三层交换机配置、路由配置、DHCP和VPN配置。', indent=True, which=doc2)
add_p('目标3（应用与安全）：能实现Web/DNS/电子邮件/FTP等服务配置及安全配置。', indent=True, which=doc2)

add_h('二、设计任务', which=doc2)
add_h('任务一：小型网络解决方案设计', level=2, which=doc2)
add_h('阶段1：需求分析与总体方案设计', level=3, which=doc2)
add_p('任务内容：', bold=True, which=doc2)
add_p('1. 根据学生专业知识背景选择一个实验题目，对相应的小型网络进行需求分析。', indent=True, which=doc2)
add_p('2. 完成该小型网络的总体方案设计，确定网络逻辑拓扑结构和所采用的网络技术。', indent=True, which=doc2)
add_p('3. 确定主要设备的性能指标，完成设备选型和经费预算（必做，含设备型号、数量、单价和总价）。', indent=True, which=doc2)
add_p('4. 绘制网络拓扑图（含设备标识和IP地址分配）。', bold=True, indent=True, which=doc2)
add_p('交付物：网络设计方案书（含需求分析、逻辑拓扑图、设备选型表、经费预算表）', bold=True, which=doc2)
add_p('学习要求：', bold=True, which=doc2)
add_p('1. 理解网络方案设计的一般流程，掌握网络方案设计的基本方法。', indent=True, which=doc2)
add_p('2. 完成网络设计方案的编写，掌握网络逻辑拓扑结构的绘制方法。', indent=True, which=doc2)
add_p('3. 掌握设备选型和经费预算的基本方法。', indent=True, which=doc2)

add_h('任务二：局域网配置与实现', level=2, which=doc2)
add_h('阶段2：网络基础搭建与配置', level=3, which=doc2)
add_p('任务内容：', bold=True, which=doc2)
add_p('1. 在模拟器中完成网络的基本搭建以及相关配置。', indent=True, which=doc2)
add_p('2. 交换机VLAN划分、三层交换机配置以及路由配置。', indent=True, which=doc2)
add_p('3. 采用DHCP进行IP自动配置并进行调试。', indent=True, which=doc2)
add_p('4. 通过VPN技术实现远程接入内部网。', indent=True, which=doc2)
add_p('交付物：网络配置文件 + 网络拓扑截图', bold=True, which=doc2)

add_h('阶段3：应用服务与安全配置', level=3, which=doc2)
add_p('任务内容：', bold=True, which=doc2)
add_p('1. 实现Web服务、DNS服务、电子邮件服务、FTP服务等应用层服务配置。', indent=True, which=doc2)
add_p('2. 进行防DHCP欺骗、防火墙等安全配置。', indent=True, which=doc2)
add_p('交付物：应用服务配置文档 + 安全配置截图', bold=True, which=doc2)

add_h('任务三：汇报答辩', level=2, which=doc2)
add_h('阶段4：答辩准备与汇报', level=3, which=doc2)
add_p('任务内容：', bold=True, which=doc2)
add_p('1. 网络整体配置检查。', indent=True, which=doc2)
add_p('2. 准备答辩PPT。', indent=True, which=doc2)
add_p('3. 小组答辩。', indent=True, which=doc2)
add_p('交付物：答辩PPT', bold=True, which=doc2)

add_h('三、进度安排', which=doc2)
sch = doc2.add_table(rows=6, cols=3); sch.style = 'Table Grid'; sch.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['天数', '阶段', '主要任务']):
    set_cell(sch.cell(0, j), h, bold=True)
for i, v in enumerate([
    ('第1天', '需求分析与方案设计', '选题、需求分析、拓扑设计、设备选型与预算'),
    ('第2天', '网络基础搭建', '交换机VLAN划分、三层交换机配置、路由配置'),
    ('第3天', 'DHCP+VPN配置', 'IP自动配置、VPN远程接入调试'),
    ('第4天', '应用服务与安全配置', 'Web/DNS/邮件/FTP配置、防火墙'),
    ('第5天', '报告撰写+答辩准备', '课程设计报告（含遇到的问题及解决方案）、答辩PPT'),
]):
    for j, v2 in enumerate(v):
        set_cell(sch.cell(i+1, j), v2)
doc2.add_paragraph()

add_h('四、思政融合要求', which=doc2)
add_p('（同初版，不变）', indent=True, which=doc2)
si = doc2.add_table(rows=5, cols=3); si.style = 'Table Grid'; si.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['阶段', '思政融入点', '融入方式']):
    set_cell(si.cell(0, j), h, bold=True)
for i, v in enumerate([
    ('需求分析', '工程标准意识：网络方案从实际出发', '讲解'),
    ('交换机/路由配置', '职业责任：配置失误可能导致全网瘫痪', '实操/检查'),
    ('DNS/应用服务', '技术自主：我国根服务器现状与科技创新', '讲解/讨论'),
    ('汇报答辩', '协作责任：模块衔接意识', '答辩'),
]):
    for j, v2 in enumerate(v):
        set_cell(si.cell(i+1, j), v2)
doc2.add_paragraph()

add_h('五、考核标准', which=doc2)
ex = doc2.add_table(rows=4, cols=3); ex.style = 'Table Grid'; ex.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['考核项目', '权重', '考查点']):
    set_cell(ex.cell(0, j), h, bold=True)
for i, v in enumerate([
    ('平时考核', '40%', '方案设计、搭建、配置、服务的完成情况'),
    ('答辩考核', '20%', '配置合理性、团队协作、问题解决能力'),
    ('设计报告考核', '40%', '正确性、完整性、规范性（含问题记录）'),
]):
    for j, v2 in enumerate(v):
        set_cell(ex.cell(i+1, j), v2)
doc2.add_paragraph()

add_h('六、交付物清单', which=doc2)
dl = doc2.add_table(rows=8, cols=3); dl.style = 'Table Grid'; dl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['序号', '交付物', '提交形式']):
    set_cell(dl.cell(0, j), h, bold=True)
for i, v in enumerate([
    ('1', '网络设计方案书（含需求分析、拓扑图、设备选型、经费预算）', 'Word / PDF'),
    ('2', '网络拓扑图（含设备标识和IP地址分配）', 'Visio / 绘图截图'),
    ('3', '网络配置文件（交换机/路由器配置命令）', '.txt / .pkt'),
    ('4', '应用服务配置截图 + 安全配置截图', 'Word / PDF'),
    ('5', '网络拓扑搭建截图（含VLAN/路由/DHCP/VPN）', 'Word / PDF'),
    ('6', '课程设计报告（含遇到的问题及解决方案）', 'Word / PDF'),
    ('7', '答辩PPT', 'PPT'),
]):
    for j, v2 in enumerate(v):
        set_cell(dl.cell(i+1, j), v2)
doc2.add_paragraph()

add_h('七、参考资料', which=doc2)
add_p('（同初版，不变）', indent=True, which=doc2)

doc2.save(os.path.join(out_dir, '示例_计网课程设计任务书_定稿.docx'))
print('DONE2')
print(f'优化报告: {os.path.join(out_dir, "示例_计网优化总结报告+对比表.docx")}')
print(f'任务书定稿: {os.path.join(out_dir, "示例_计网课程设计任务书_定稿.docx")}')
