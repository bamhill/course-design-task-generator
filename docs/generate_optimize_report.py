# -*- coding: utf-8 -*-
"""生成优化总结报告 + 改动对比表 WORD 版"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def set_cell(cell, text, bold=False, size=Pt(10.5)):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = size; run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_p(text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ═══ 标题 ═══
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('课程设计任务书优化总结报告')
r.font.name = '黑体'; r.font.size = Pt(18); r.bold = True
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('《数据库原理与应用（SQL-Server）实践》')
r2.font.name = '宋体'; r2.font.size = Pt(12)
r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()

# ═══ 一、优化概览 ═══
add_h('一、优化概览')
add_p('本报告记录了课程设计任务书初版在教师审阅后的优化过程。教师共提出 2 条优化需求，涉及进度安排调整和考核权重调整。以下逐条说明优化内容、原因及结果。', indent=True)

# ═══ 二、优化条目 ═══
add_h('二、优化条目详情', level=1)

# 条目1
add_h('优化条目 #1：概念设计阶段时间调整', level=2)
add_p('原内容：概念结构设计仅分配1天时间（第2天），需在同一天完成局部E-R图设计、合并、冲突解决等全部工作。', indent=True)
add_p('优化后：概念结构设计拆分为2天——第2天做局部E-R图设计，新增检查节点；第3天做E-R图合并+冲突解决，与逻辑结构设计衔接。', indent=True)
add_p('优化原因：教师审阅后认为概念设计阶段是数据库设计的核心环节，任务量大，1天时间过于紧张，学生容易草率完成，影响后续所有阶段的质量。', indent=True)
add_p('涉及章节：§三 进度安排', indent=True)

# 条目2
add_h('优化条目 #2：考核权重调整', level=2)
add_p('原内容：平时考核20%，设计报告考核60%，答辩考核20%。', indent=True)
add_p('优化后：平时考核30%，设计报告考核50%，答辩考核20%。', indent=True)
add_p('优化原因：教师认为过程管理比结果更重要，提高平时考核权重可以更好地督促学生按进度完成各阶段任务，降低期末突击风险。', indent=True)
add_p('涉及章节：§五 考核标准', indent=True)

doc.add_paragraph()

# ═══ 三、改动对比表 ═══
add_h('三、改动对比表', level=1)

comp_table = doc.add_table(rows=3, cols=4)
comp_table.style = 'Table Grid'
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

comp_headers = ['#', '改动位置', '原内容', '改后内容']
comp_data = [
    ('1', '§三 进度安排', '概念结构设计 1天（第2天）', '概念结构设计 2天（第2-3天，含检查节点）'),
    ('2', '§五 考核标准', '平时20% + 报告60% + 答辩20%', '平时30% + 报告50% + 答辩20%'),
]
for j, h in enumerate(comp_headers):
    set_cell(comp_table.cell(0, j), h, bold=True)
for i, row in enumerate(comp_data):
    for j, v in enumerate(row):
        set_cell(comp_table.cell(i+1, j), v)

doc.add_paragraph()

# ═══ 四、优化结果 ═══
add_h('四、优化结果')
add_p('以上 2 条优化需求已全部融入任务书定稿。', indent=True)
add_p('• 进度安排表已更新：第2天→局部E-R图设计，第3天→E-R图合并+逻辑结构设计', indent=True)
add_p('• 考核标准已更新：平时30% / 报告50% / 答辩20%', indent=True)
add_p('• 其余章节内容不变，与原初版保持一致', indent=True)

doc.add_paragraph()

# ═══ 最终版标记 ═══
add_h('附：课程设计任务书（定稿）')
add_p('优化后的任务书定稿文件为：', indent=True)
add_p('《数据库原理与应用（SQL-Server）实践》课程设计任务书（定稿）.docx', indent=True)
add_p('该文件已包含以上全部优化内容，可直接用于下发学生。', indent=True)

# 保存
output_dir = os.path.dirname(os.path.abspath(__file__))
doc.save(os.path.join(output_dir, '示例_优化总结报告+对比表.docx'))
print(f'DONE: {os.path.join(output_dir, "示例_优化总结报告+对比表.docx")}')
